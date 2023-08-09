from typing import Iterable, Sequence

from document.config import settings
from document.domain.assembly_strategies.assembly_strategy_utils import (
    adjust_book_intro_headings,
    book_intro_commentary,
    chapter_commentary,
    chapter_intro,
    ensure_primary_usfm_books_for_different_languages_are_adjacent,
    verses_for_chapter_tn,
    verses_for_chapter_tq,
)
from document.domain.assembly_strategies_docx.assembly_strategy_utils import (
    add_one_column_section,
    add_two_column_section,
    add_hr,
    add_page_break,
    create_docx_subdoc,
)
from document.domain.model import (
    BCBook,
    BookContent,
    HtmlContent,
    LangDirEnum,
    TNBook,
    TQBook,
    TWBook,
    USFMBook,
)
from document.utils.number_utils import is_even
from docx import Document  # type: ignore
from docxtpl import DocxTemplate  # type: ignore
from htmldocx import HtmlToDocx  # type: ignore
from docx.enum.section import WD_SECTION  # type: ignore
from docx.enum.text import WD_BREAK  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docxcompose.composer import Composer  # type: ignore

logger = settings.logger(__name__)


def assemble_usfm_as_iterator_by_chapter_for_book_then_lang_1c(
    usfm_book_content_units: Sequence[USFMBook],
    tn_book_content_units: Sequence[TNBook],
    tq_book_content_units: Sequence[TQBook],
    tw_book_content_units: Sequence[TWBook],
    bc_book_content_units: Sequence[BCBook],
    footnotes_heading: HtmlContent = settings.FOOTNOTES_HEADING,
    tn_verse_notes_enclosing_div_fmt_str: str = settings.TN_VERSE_NOTES_ENCLOSING_DIV_FMT_STR,
    tq_heading_and_questions_fmt_str: str = settings.TQ_HEADING_AND_QUESTIONS_FMT_STR,
) -> Composer:
    """
    Construct the Docx wherein at least one USFM resource exists, one column
    layout.
    """

    ldebug = logger.debug
    lexception = logger.exception

    # Sort resources by language
    def sort_key(resource: BookContent) -> str:
        return resource.lang_code

    usfm_book_content_units = sorted(usfm_book_content_units, key=sort_key)
    tn_book_content_units = sorted(tn_book_content_units, key=sort_key)
    tq_book_content_units = sorted(tq_book_content_units, key=sort_key)
    bc_book_content_units = sorted(bc_book_content_units, key=sort_key)

    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    # Add book intros for each tn_book_content_unit
    for tn_book_content_unit in tn_book_content_units:
        if tn_book_content_unit.intro_html:
            book_intro = tn_book_content_unit.intro_html
            book_intro = adjust_book_intro_headings(book_intro)

            subdoc = create_docx_subdoc(
                book_intro,
                tn_book_content_unit.lang_code,
                tn_book_content_unit
                and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)

    # Add book commentary for each bc_book_content_units
    for bc_book_content_unit in bc_book_content_units:
        subdoc = create_docx_subdoc(
            "".join(bc_book_content_unit.book_intro), bc_book_content_unit.lang_code
        )
        composer.append(subdoc)

    # Use the usfm_book_content_unit that has the most chapters as a
    # chapter_num pump to realize the most amount of content displayed
    # to the user.
    usfm_with_most_chapters = max(
        usfm_book_content_units,
        key=lambda usfm_book_content_unit: usfm_book_content_unit.chapters.keys(),
    )
    for chapter_num, chapter in usfm_with_most_chapters.chapters.items():
        add_one_column_section(doc)

        chapter_intro_ = HtmlContent("")
        chapter_commentary_ = HtmlContent("")

        # Add chapter intro for each language
        for tn_book_content_unit2 in tn_book_content_units:
            # NOTE For v1 we aren't yet decided about where or whether
            # we'll interleave chapter intros.
            # Add the translation notes chapter intro.
            chapter_intro_ = chapter_intro(tn_book_content_unit2, chapter_num)
            if chapter_intro_:
                subdoc = create_docx_subdoc(
                    chapter_intro_,
                    tn_book_content_unit2.lang_code,
                    tn_book_content_unit2
                    and tn_book_content_unit2.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)

        for bc_book_content_unit in bc_book_content_units:
            # NOTE For v1 we aren't yet decided about where or whether
            # we'll interleave chapter commentary.
            # Add the chapter commentary.
            chapter_commentary_ = chapter_commentary(bc_book_content_unit, chapter_num)
            subdoc = create_docx_subdoc(
                chapter_commentary_, bc_book_content_unit.lang_code
            )
            composer.append(subdoc)

        # Add the interleaved USFM chapters
        for usfm_book_content_unit in usfm_book_content_units:
            if chapter_num in usfm_book_content_unit.chapters:
                add_one_column_section(doc)

                # fmt: off
                is_rtl = usfm_book_content_unit and usfm_book_content_unit.lang_direction == LangDirEnum.RTL
                # fmt: on

                subdoc = create_docx_subdoc(
                    "".join(usfm_book_content_unit.chapters[chapter_num].content),
                    usfm_book_content_unit.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)

                try:
                    chapter_footnotes = usfm_book_content_unit.chapters[
                        chapter_num
                    ].footnotes
                    if chapter_footnotes:
                        add_one_column_section(doc)

                        subdoc = create_docx_subdoc(
                            "{}{}".format(footnotes_heading, chapter_footnotes),
                            usfm_book_content_unit.lang_code,
                            is_rtl,
                        )
                        composer.append(subdoc)

                except KeyError:
                    ldebug(
                        "usfm_book_content_unit: %s, does not have chapter: %s",
                        usfm_book_content_unit,
                        chapter_num,
                    )
                    lexception("Caught exception:")

        # Add the interleaved tn notes
        tn_verses = None
        for tn_book_content_unit3 in tn_book_content_units:
            tn_verses = verses_for_chapter_tn(tn_book_content_unit3, chapter_num)
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tn_verse_notes_enclosing_div_fmt_str.format(
                        "".join(tn_verses.values())
                    ),
                    tn_book_content_unit3.lang_code,
                    tn_book_content_unit3
                    and tn_book_content_unit3.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)

        # Add the interleaved tq questions
        tq_verses = None
        for tq_book_content_unit in tq_book_content_units:
            tq_verses = verses_for_chapter_tq(tq_book_content_unit, chapter_num)
            # Add TQ verse content, if any
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_heading_and_questions_fmt_str.format(
                        tq_book_content_unit.resource_type_name,
                        "".join(tq_verses.values()),
                    ),
                    tq_book_content_unit.lang_code,
                    tq_book_content_unit
                    and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)

        add_page_break(doc)

    return composer


def assemble_tn_as_iterator_by_chapter_for_book_then_lang(
    usfm_book_content_units: Sequence[USFMBook],
    tn_book_content_units: Sequence[TNBook],
    tq_book_content_units: Sequence[TQBook],
    tw_book_content_units: Sequence[TWBook],
    bc_book_content_units: Sequence[BCBook],
    tn_verse_notes_enclosing_div_fmt_str: str = settings.TN_VERSE_NOTES_ENCLOSING_DIV_FMT_STR,
    tq_heading_and_questions_fmt_str: str = settings.TQ_HEADING_AND_QUESTIONS_FMT_STR,
) -> Composer:
    """
    Construct the HTML for a 'by chapter' strategy wherein at least
    tn_book_content_units exists.
    """
    # Sort resources by language
    def sort_key(resource: BookContent) -> str:
        return resource.lang_code

    tn_book_content_units = sorted(tn_book_content_units, key=sort_key)
    tq_book_content_units = sorted(tq_book_content_units, key=sort_key)
    bc_book_content_units = sorted(bc_book_content_units, key=sort_key)

    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    add_one_column_section(doc)

    for tn_book_content_unit in tn_book_content_units:
        if tn_book_content_unit.intro_html:
            # Add the book intro
            book_intro = tn_book_content_unit.intro_html
            book_intro = adjust_book_intro_headings(book_intro)
            subdoc = create_docx_subdoc(
                book_intro,
                tn_book_content_unit.lang_code,
                tn_book_content_unit
                and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)

    for bc_book_content_unit in bc_book_content_units:
        subdoc = create_docx_subdoc(
            book_intro_commentary(bc_book_content_unit), bc_book_content_unit.lang_code
        )
        composer.append(subdoc)

    # Use the tn_book_content_unit that has the most chapters as a
    # chapter_num pump to realize the most amount of content displayed
    # to user.
    tn_with_most_chapters = max(
        tn_book_content_units,
        key=lambda tn_book_content_unit: tn_book_content_unit.chapters.keys(),
    )
    for chapter_num in tn_with_most_chapters.chapters.keys():
        add_one_column_section(doc)
        one_column_html = []
        one_column_html.append("Chapter {}".format(chapter_num))

        for tn_book_content_unit in tn_book_content_units:

            # Add the translation notes chapter intro.
            one_column_html.append(chapter_intro(tn_book_content_unit, chapter_num))
            one_column_html_ = "".join(one_column_html)
            if one_column_html_:
                subdoc = create_docx_subdoc(
                    one_column_html_,
                    tn_book_content_unit.lang_code,
                    tn_book_content_unit
                    and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)

        for bc_book_content_unit in bc_book_content_units:
            # Add the chapter commentary.
            subdoc = create_docx_subdoc(
                chapter_commentary(bc_book_content_unit, chapter_num),
                bc_book_content_unit.lang_code,
            )
            composer.append(subdoc)

        # Add the interleaved tn notes
        for tn_book_content_unit in tn_book_content_units:
            tn_verses = verses_for_chapter_tn(tn_book_content_unit, chapter_num)
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tn_verse_notes_enclosing_div_fmt_str.format(
                        "".join(tn_verses.values())
                    ),
                    tn_book_content_unit.lang_code,
                    tn_book_content_unit
                    and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)

        # Add the interleaved tq questions
        for tq_book_content_unit in tq_book_content_units:
            tq_verses = verses_for_chapter_tq(tq_book_content_unit, chapter_num)
            # Add TQ verse content, if any
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_heading_and_questions_fmt_str.format(
                        tq_book_content_unit.resource_type_name,
                        "".join(tq_verses.values()),
                    ),
                    tq_book_content_unit.lang_code,
                    tq_book_content_unit
                    and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
        add_page_break(doc)
    return composer


def assemble_tq_as_iterator_by_chapter_for_book_then_lang(
    usfm_book_content_units: Sequence[USFMBook],
    tn_book_content_units: Sequence[TNBook],
    tq_book_content_units: Sequence[TQBook],
    tw_book_content_units: Sequence[TWBook],
    bc_book_content_units: Sequence[BCBook],
    tq_heading_and_questions_fmt_str: str = settings.TQ_HEADING_AND_QUESTIONS_FMT_STR,
) -> Composer:
    """
    Construct the HTML for a 'by chapter' strategy wherein at least
    tq_book_content_units exists.
    """
    # Sort resources by language
    def sort_key(resource: BookContent) -> str:
        return resource.lang_code

    tq_book_content_units = sorted(tq_book_content_units, key=sort_key)
    bc_book_content_units = sorted(bc_book_content_units, key=sort_key)

    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    # Use the tq_book_content_unit that has the most chapters as a
    # chapter_num pump to realize the most amount of content displayed to user.
    tq_with_most_chapters = max(
        tq_book_content_units,
        key=lambda tq_book_content_unit: tq_book_content_unit.chapters.keys(),
    )
    for chapter_num in tq_with_most_chapters.chapters.keys():
        one_column_html = []
        one_column_html.append("Chapter {}".format(chapter_num))

        for bc_book_content_unit in bc_book_content_units:
            # Add the chapter commentary.
            one_column_html.append(
                chapter_commentary(bc_book_content_unit, chapter_num)
            )

        if one_column_html:
            add_one_column_section(doc)

            subdoc = create_docx_subdoc(
                "".join(one_column_html), tq_with_most_chapters.lang_code
            )
            composer.append(subdoc)

        # Add the interleaved tq questions
        for tq_book_content_unit in tq_book_content_units:
            tq_verses = verses_for_chapter_tq(tq_book_content_unit, chapter_num)
            # Add TQ verse content, if any
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    tq_heading_and_questions_fmt_str.format(
                        tq_book_content_unit.resource_type_name,
                        "".join(tq_verses.values()),
                    ),
                    tq_book_content_unit.lang_code,
                    tq_book_content_unit
                    and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
        add_page_break(doc)
    return composer


def assemble_tw_as_iterator_by_chapter_for_book_then_lang(
    usfm_book_content_units: Sequence[USFMBook],
    tn_book_content_units: Sequence[TNBook],
    tq_book_content_units: Sequence[TQBook],
    tw_book_content_units: Sequence[TWBook],
    bc_book_content_units: Sequence[BCBook],
) -> Composer:
    """Construct the HTML for BC and TW."""

    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    # Sort resources by language
    def sort_key(resource: BookContent) -> str:
        return resource.lang_code

    bc_book_content_units = sorted(bc_book_content_units, key=sort_key)

    # Add the bible commentary
    for bc_book_content_unit in bc_book_content_units:
        subdoc = create_docx_subdoc(
            bc_book_content_unit.book_intro, bc_book_content_unit.lang_code
        )
        composer.append(subdoc)
        for chapter in bc_book_content_unit.chapters.values():
            subdoc = create_docx_subdoc(
                chapter.commentary, bc_book_content_unit.lang_code
            )
            composer.append(subdoc)

            add_page_break(doc)

    return composer
