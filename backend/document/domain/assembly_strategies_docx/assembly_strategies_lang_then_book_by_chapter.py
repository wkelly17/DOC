from typing import Optional

from document.config import settings
from document.domain.assembly_strategies.assembly_strategy_utils import (
    # commentary_book_intro,
    book_title,
    chapter_commentary,
    chapter_content_sans_footnotes,
    chapter_heading,
    chapter_intro,
    tn_chapter_verses,
    tq_chapter_verses,
)
from document.domain.assembly_strategies_docx.assembly_strategy_utils import (
    add_hr,
    create_docx_subdoc,
    add_one_column_section,
    add_two_column_section,
    add_page_break,
)

from document.domain.model import (
    BCBook,
    HtmlContent,
    LangDirEnum,
    TNBook,
    TQBook,
    TWBook,
    USFMBook,
)
from docx import Document  # type: ignore
from htmldocx import HtmlToDocx  # type: ignore
from docxcompose.composer import Composer  # type: ignore

logger = settings.logger(__name__)


def assemble_usfm_by_book(
    usfm_book_content_unit: Optional[USFMBook],
    tn_book_content_unit: Optional[TNBook],
    tq_book_content_unit: Optional[TQBook],
    tw_book_content_unit: Optional[TWBook],
    usfm_book_content_unit2: Optional[USFMBook],
    bc_book_content_unit: Optional[BCBook],
    footnotes_heading: HtmlContent = settings.FOOTNOTES_HEADING,
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    usfm_book_content_unit exists.
    """
    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    if tn_book_content_unit:
        if tn_book_content_unit.book_intro:
            subdoc = create_docx_subdoc(
                tn_book_content_unit.book_intro,
                tn_book_content_unit.lang_code,
                tn_book_content_unit
                and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
    if bc_book_content_unit:
        if bc_book_content_unit.book_intro:
            subdoc = create_docx_subdoc(
                "".join(bc_book_content_unit.book_intro),
                bc_book_content_unit.lang_code,
            )
            composer.append(subdoc)
    if usfm_book_content_unit:
        # fmt: off
        is_rtl = usfm_book_content_unit and usfm_book_content_unit.lang_direction == LangDirEnum.RTL
        # fmt: on
        subdoc = create_docx_subdoc(
            "".join(book_title(usfm_book_content_unit.book_code)),
            usfm_book_content_unit.lang_code,
            is_rtl,
        )
        composer.append(subdoc)
        for (
            chapter_num,
            chapter,
        ) in usfm_book_content_unit.chapters.items():
            add_one_column_section(doc)
            tn_verses: list[HtmlContent] = []
            tq_verses: list[HtmlContent] = []
            chapter_intro_ = HtmlContent("")
            chapter_commentary_ = HtmlContent("")
            if tn_book_content_unit:
                chapter_intro_ = "".join(
                    chapter_intro(tn_book_content_unit, chapter_num)
                )
                tn_verses = list(tn_chapter_verses(tn_book_content_unit, chapter_num))
            if bc_book_content_unit:
                chapter_commentary_ = "".join(
                    chapter_commentary(bc_book_content_unit, chapter_num)
                )
            if tq_book_content_unit:
                tq_verses = list(tq_chapter_verses(tq_book_content_unit, chapter_num))
            # Footnotes are rendered to HTML by USFM-TOOLS after the chapter's
            # verse content so we don't want them included additionally/accidentally
            # here also since we explicitly include footnotes after this. Later, if
            # we ditch the verse level interleaving, we can move this into the
            # parsing module.
            #
            # Find the index of where footnotes occur at the end of
            # chapter.content, if the chapter has footnotes.
            subdoc = create_docx_subdoc(
                "".join(chapter_content_sans_footnotes(chapter.content)),
                usfm_book_content_unit.lang_code,
                is_rtl,
            )
            composer.append(subdoc)
            # Add scripture footnotes if available
            if chapter.footnotes:
                subdoc = create_docx_subdoc(
                    "{}{}".format(footnotes_heading, chapter.footnotes),
                    usfm_book_content_unit.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)
            if chapter_intro_:
                subdoc = create_docx_subdoc(
                    chapter_intro_, usfm_book_content_unit.lang_code, is_rtl
                )
                composer.append(subdoc)
            if chapter_commentary_:
                subdoc = create_docx_subdoc(
                    chapter_commentary_, usfm_book_content_unit.lang_code, is_rtl
                )
                composer.append(subdoc)
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    "".join(tn_verses),
                    usfm_book_content_unit.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    "".join(tq_verses),
                    usfm_book_content_unit.lang_code,
                    is_rtl,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            # TODO Get feedback on whether we should allow a user to select a primary _and_
            # a secondary USFM resource. If we want to limit the user to only one USFM per
            # document then we would want to control that in the UI and maybe also at the API
            # level. The API level control would be implemented in the DocumentRequest
            # validation.
            if usfm_book_content_unit2:
                add_one_column_section(doc)
                # Here we add the whole chapter's worth of verses for the secondary usfm
                subdoc = create_docx_subdoc(
                    "".join(usfm_book_content_unit2.chapters[chapter_num].content),
                    usfm_book_content_unit.lang_code,
                    usfm_book_content_unit2
                    and usfm_book_content_unit2.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tn_by_book(
    usfm_book_content_unit: Optional[USFMBook],
    tn_book_content_unit: Optional[TNBook],
    tq_book_content_unit: Optional[TQBook],
    tw_book_content_unit: Optional[TWBook],
    usfm_book_content_unit2: Optional[USFMBook],
    bc_book_content_unit: Optional[BCBook],
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    tn_book_content_unit exists.
    """
    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    if tn_book_content_unit:
        if tn_book_content_unit.book_intro:
            # Add the chapter intro
            subdoc = create_docx_subdoc(
                "".join(tn_book_content_unit.book_intro),
                tn_book_content_unit.lang_code,
                tn_book_content_unit
                and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
        if bc_book_content_unit:
            if bc_book_content_unit.book_intro:
                subdoc = create_docx_subdoc(
                    "".join(bc_book_content_unit.book_intro),
                    tn_book_content_unit.lang_code,
                )
                composer.append(subdoc)
        for chapter_num in tn_book_content_unit.chapters:
            add_one_column_section(doc)
            one_column_html = []
            one_column_html.append(chapter_heading(tn_book_content_unit, chapter_num))
            one_column_html.append(
                "".join(chapter_intro(tn_book_content_unit, chapter_num))
            )
            one_column_html_ = "".join(one_column_html)
            if one_column_html_:
                subdoc = create_docx_subdoc(
                    one_column_html_,
                    tn_book_content_unit.lang_code,
                    tn_book_content_unit
                    and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            if bc_book_content_unit:
                subdoc = create_docx_subdoc(
                    "".join(chapter_commentary(bc_book_content_unit, chapter_num)),
                    bc_book_content_unit.lang_code,
                )
                composer.append(subdoc)
            tn_verses = list(tn_chapter_verses(tn_book_content_unit, chapter_num))
            if tn_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    "".join(tn_verses),
                    tn_book_content_unit.lang_code,
                    tn_book_content_unit
                    and tn_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            tq_verses = list(tq_chapter_verses(tq_book_content_unit, chapter_num))
            if tq_book_content_unit and tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    "".join(tq_verses),
                    tq_book_content_unit.lang_code,
                    tq_book_content_unit
                    and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
                add_one_column_section(doc)
                p = doc.add_paragraph()
                add_hr(p)
            add_page_break(doc)
    return composer


def assemble_tq_by_book(
    usfm_book_content_unit: Optional[USFMBook],
    tn_book_content_unit: Optional[TNBook],
    tq_book_content_unit: Optional[TQBook],
    tw_book_content_unit: Optional[TWBook],
    usfm_book_content_unit2: Optional[USFMBook],
    bc_book_content_unit: Optional[BCBook],
) -> Composer:
    """
    Construct the HTML for a 'by book' strategy wherein at least
    tq_book_content_unit exists.
    """
    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    if tq_book_content_unit:
        for chapter_num in tq_book_content_unit.chapters:
            add_one_column_section(doc)
            if bc_book_content_unit:
                subdoc = create_docx_subdoc(
                    "".join(chapter_commentary(bc_book_content_unit, chapter_num)),
                    bc_book_content_unit.lang_code,
                )
                composer.append(subdoc)
            subdoc = create_docx_subdoc(
                chapter_heading(tq_book_content_unit, chapter_num),
                tq_book_content_unit.lang_code,
                tq_book_content_unit
                and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
            )
            composer.append(subdoc)
            tq_verses = list(tq_chapter_verses(tq_book_content_unit, chapter_num))
            if tq_verses:
                add_two_column_section(doc)
                subdoc = create_docx_subdoc(
                    "".join(tq_verses),
                    tq_book_content_unit.lang_code,
                    tq_book_content_unit
                    and tq_book_content_unit.lang_direction == LangDirEnum.RTL,
                )
                composer.append(subdoc)
            add_page_break(doc)
    return composer


def assemble_tw_by_book(
    usfm_book_content_unit: Optional[USFMBook],
    tn_book_content_unit: Optional[TNBook],
    tq_book_content_unit: Optional[TQBook],
    tw_book_content_unit: Optional[TWBook],
    usfm_book_content_unit2: Optional[USFMBook],
    bc_book_content_unit: Optional[BCBook],
) -> Composer:

    """
    TW is handled outside this module, that is why no
    code for TW is explicitly included here.
    """
    html_to_docx = HtmlToDocx()
    doc = Document()
    composer = Composer(doc)

    if bc_book_content_unit:
        subdoc = create_docx_subdoc(
            bc_book_content_unit.book_intro, bc_book_content_unit.lang_code
        )
        composer.append(subdoc)
        for chapter in bc_book_content_unit.chapters.values():
            subdoc = create_docx_subdoc(
                chapter.commentary, bc_book_content_unit.lang_code
            )
            composer.append(subdoc)
            add_page_break(doc)
    return composer
